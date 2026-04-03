from pathlib import Path
from docx import Document
from src.utils.text_utils import normalize_whitespace


class DOCXExtractor:
    def extract(self, file_path: Path) -> tuple[str, bool]:
        document = Document(file_path)
        blocks = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                blocks.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        return normalize_whitespace("\n".join(blocks)), False